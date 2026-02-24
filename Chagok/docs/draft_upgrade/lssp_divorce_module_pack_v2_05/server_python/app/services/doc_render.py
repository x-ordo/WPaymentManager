from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Dict, List, Literal, Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DocType = Literal[
    "CONSENSUAL_INTENT_FORM",
    "CHILD_CUSTODY_AGREEMENT",
    "EVIDENCE_INDEX",
    "ASSET_INVENTORY",
    "CASE_TIMELINE_SUMMARY",
]

def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _add_kv_table(doc: Document, rows: List[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = str(k)
        table.cell(i, 1).text = str(v)

def render_docx(doc_type: DocType, payload: Dict[str, Any], out_path: Path) -> Path:
    doc = Document()

    if doc_type == "CONSENSUAL_INTENT_FORM":
        _add_title(doc, "협의이혼의사확인신청서 (자동생성)")
        doc.add_paragraph("※ 제출 전 반드시 사실관계를 확인하세요.")
        p = payload["parties"]["petitioner"]
        r = payload["parties"]["respondent"]
        doc.add_paragraph("1) 당사자")
        _add_kv_table(doc, [
            ("신청인(배우자 1) 성명", p["name"]),
            ("주민번호(마스킹)", p["rrn_masked"]),
            ("주소", p["address"]),
            ("연락처", p["phone"]),
            ("상대방(배우자 2) 성명", r["name"]),
            ("주민번호(마스킹)", r["rrn_masked"]),
            ("주소", r["address"]),
            ("연락처", r["phone"]),
        ])
        ws = payload["parties"]["witnesses"]
        doc.add_paragraph("2) 증인(성년 2명)")
        for idx, w in enumerate(ws, start=1):
            doc.add_paragraph(f"- 증인 {idx}")
            _add_kv_table(doc, [
                ("성명", w["name"]),
                ("주민번호(마스킹)", w["rrn_masked"]),
                ("주소", w["address"]),
                ("연락처", w["phone"]),
                ("서명/날인", "________________"),
            ])
        doc.add_paragraph("작성일: __________________")
        doc.add_paragraph("신청인(배우자 1) 서명: __________________")
        doc.add_paragraph("신청인(배우자 2) 서명: __________________")

    elif doc_type == "CHILD_CUSTODY_AGREEMENT":
        _add_title(doc, "자녀 양육·친권자 결정 협의서 (자동생성)")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "성명"
        hdr[1].text = "생년월일"
        hdr[2].text = "주 양육자"
        hdr[3].text = "비고"
        for c in payload.get("children", []):
            row = table.add_row().cells
            row[0].text = c["name"]
            row[1].text = c["birth"]
            row[2].text = "신청인" if c["custody_parent"] == "petitioner" else "상대방"
            row[3].text = c.get("notes", "")

        terms = payload["custody_terms"]
        doc.add_paragraph("2) 양육·친권 및 면접교섭")
        _add_kv_table(doc, [
            ("친권 및 주 양육자", "신청인" if terms["custody"] == "petitioner" else "상대방"),
            ("면접교섭", terms["visitation"]),
        ])
        doc.add_paragraph("3) 양육비")
        _add_kv_table(doc, [
            ("월 양육비(원)", f'{terms["child_support_monthly"]:,}'),
            ("지급일", terms["payment_day"]),
            ("특수비용", terms.get("special_costs", "")),
        ])
        doc.add_paragraph("작성일: __________________")
        doc.add_paragraph("신청인 서명: __________________")
        doc.add_paragraph("상대방 서명: __________________")

    elif doc_type == "EVIDENCE_INDEX":
        _add_title(doc, "증거목록표 (갑1~) (자동생성)")
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "번호"
        hdr[1].text = "종류"
        hdr[2].text = "요지"
        hdr[3].text = "발생/작성일"
        hdr[4].text = "연결근거"
        hdr[5].text = "상태"
        for e in payload.get("evidences", []):
            row = table.add_row().cells
            row[0].text = e["no"]
            row[1].text = e["type"]
            row[2].text = e["summary"]
            row[3].text = e.get("occurred_at", "")
            row[4].text = ",".join(e.get("grounds", []))
            row[5].text = e.get("status", "")

    elif doc_type == "ASSET_INVENTORY":
        _add_title(doc, "재산목록표 (자동생성)")
        doc.add_paragraph("1) 자산")
        t1 = doc.add_table(rows=1, cols=5)
        t1.style = "Table Grid"
        h = t1.rows[0].cells
        h[0].text="구분"; h[1].text="명칭"; h[2].text="가액(원)"; h[3].text="명의"; h[4].text="비고"
        for a in payload.get("assets", []):
            r = t1.add_row().cells
            r[0].text=a["type"]; r[1].text=a["name"]; r[2].text=f'{a["value"]:,}'; r[3].text=a.get("owner",""); r[4].text=a.get("notes","")
        doc.add_paragraph("2) 채무")
        t2 = doc.add_table(rows=1, cols=5)
        t2.style = "Table Grid"
        h = t2.rows[0].cells
        h[0].text="구분"; h[1].text="명칭"; h[2].text="금액(원)"; h[3].text="채무자"; h[4].text="비고"
        for d in payload.get("debts", []):
            r = t2.add_row().cells
            r[0].text=d["type"]; r[1].text=d["name"]; r[2].text=f'{d["value"]:,}'; r[3].text=d.get("debtor",""); r[4].text=d.get("notes","")
        doc.add_paragraph("작성일: __________________")
        doc.add_paragraph("제출인 서명: __________________")

    elif doc_type == "CASE_TIMELINE_SUMMARY":
        _add_title(doc, "사실관계 요지서(타임라인) (자동생성)")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text="일자"; hdr[1].text="사건"; hdr[2].text="유형"
        for ev in payload.get("timeline", []):
            row = table.add_row().cells
            row[0].text = ev["date"]
            row[1].text = ev["event"]
            row[2].text = ev["type"]

    else:
        raise ValueError(f"unsupported doc_type: {doc_type}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
