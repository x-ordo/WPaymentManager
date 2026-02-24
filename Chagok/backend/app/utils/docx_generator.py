"""
DocxGenerator - Korean Legal Document DOCX Generation Utility

Generates Word documents with proper Korean legal formatting:
- A4 paper size (210mm x 297mm)
- Korean court margins (25mm top/bottom, 20mm left/right)
- Batang/Noto Serif CJK KR font at 12pt
- Line spacing 1.6
- Proper section structure for legal documents
"""

from io import BytesIO
from typing import Dict, Any, List, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxGeneratorError(Exception):
    """Custom exception for DOCX generation errors"""
    pass


class DocxGenerator:
    """
    Generator for Korean legal documents in DOCX format.

    Follows Korean court document standards:
    - A4 paper size
    - 25mm top/bottom margins, 20mm left/right margins
    - Batang or Noto Serif CJK KR font
    - 12pt body text, 1.6 line spacing
    """

    # Korean legal document constants
    PAPER_WIDTH_MM = 210
    PAPER_HEIGHT_MM = 297
    MARGIN_TOP_MM = 25
    MARGIN_BOTTOM_MM = 25
    MARGIN_LEFT_MM = 20
    MARGIN_RIGHT_MM = 20

    # Font settings
    FONT_NAME_PRIMARY = "바탕"  # Batang
    FONT_NAME_FALLBACK = "Batang"
    FONT_NAME_MODERN = "Noto Serif CJK KR"
    FONT_SIZE_BODY = 12
    FONT_SIZE_TITLE = 18
    FONT_SIZE_HEADING1 = 14
    FONT_SIZE_HEADING2 = 12
    LINE_SPACING = 1.6

    # Document type titles in Korean
    DOCUMENT_TYPES = {
        "complaint": "소 장",
        "motion": "신 청 서",
        "brief": "준 비 서 면",
        "response": "답 변 서",
    }

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise DocxGeneratorError(
                "python-docx is not installed. "
                "Please install it with: pip install python-docx"
            )
        self.doc = None

    def generate_document(
        self,
        content: Dict[str, Any],
        document_type: str = "brief"
    ) -> bytes:
        """
        Generate a DOCX document from structured content.

        Args:
            content: Dictionary with structure:
                {
                    "header": {
                        "court_name": "서울가정법원",
                        "case_number": "2025드단12345",
                        "parties": {
                            "plaintiff": "원고 김○○",
                            "defendant": "피고 이○○"
                        }
                    },
                    "sections": [
                        {"title": "청구취지", "content": "...", "order": 1},
                        {"title": "청구원인", "content": "...", "order": 2}
                    ],
                    "citations": [
                        {"reference": "[증 제1호증]", "description": "..."}
                    ],
                    "footer": {
                        "date": "2025년 12월 3일",
                        "attorney": "변호사 박○○"
                    }
                }
            document_type: One of "complaint", "motion", "brief", "response"

        Returns:
            bytes: DOCX file content
        """
        self.doc = Document()
        self._setup_page_format()
        self._setup_styles()

        # Build document structure
        self._add_case_header(content.get("header", {}))
        self._add_document_title(document_type)
        self._add_sections(content.get("sections", []))
        self._add_citations(content.get("citations", []))
        self._add_signature(content.get("footer", {}))
        self._add_court_submission(content.get("header", {}).get("court_name", ""))

        # Save to bytes
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_from_lines(
        self,
        lines: List[Dict[str, Any]],
        case_title: Optional[str] = None
    ) -> bytes:
        """
        Generate DOCX from line-based template with precise formatting.

        Args:
            lines: List of line dictionaries with structure:
                {
                    "line": int,          # Line number
                    "text": str,          # Text content
                    "section": str,       # Optional section name
                    "format": {
                        "align": "left|center|right",
                        "indent": int,    # Characters to indent
                        "bold": bool,
                        "font_size": int,
                        "spacing_before": int,
                        "spacing_after": int
                    }
                }
            case_title: Optional case title for document header

        Returns:
            bytes: DOCX file content
        """
        self.doc = Document()
        self._setup_page_format()
        self._setup_styles()

        # Add case title if provided
        if case_title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"사건: {case_title}")
            run.font.size = Pt(10)
            self._set_run_korean_font(run)
            self.doc.add_paragraph()

        # Process each line
        for line_data in lines:
            text = line_data.get("text", "")
            fmt = line_data.get("format", {})

            # Create paragraph
            p = self.doc.add_paragraph()

            # Apply alignment
            align = fmt.get("align", "left")
            if align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Apply indent (convert character count to approximate points)
            indent = fmt.get("indent", 0)
            if indent > 0:
                # Approximate: 1 character = 12pt width (for Korean monospace)
                p.paragraph_format.left_indent = Pt(indent * 6)

            # Apply spacing
            spacing_before = fmt.get("spacing_before", 0)
            spacing_after = fmt.get("spacing_after", 0)
            if spacing_before > 0:
                p.paragraph_format.space_before = Pt(spacing_before * 12)
            if spacing_after > 0:
                p.paragraph_format.space_after = Pt(spacing_after * 12)

            # Add text with formatting
            if text:
                run = p.add_run(text)

                # Apply bold
                if fmt.get("bold", False):
                    run.bold = True

                # Apply font size
                font_size = fmt.get("font_size", self.FONT_SIZE_BODY)
                run.font.size = Pt(font_size)

                # Set Korean font
                self._set_run_korean_font(run)

        # Add disclaimer at the end
        self.doc.add_paragraph()
        disclaimer = self.doc.add_paragraph(
            "※ 본 문서는 AI가 생성한 초안이며, "
            "변호사의 검토 및 수정이 필요합니다."
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in disclaimer.runs:
            run.italic = True
            run.font.size = Pt(9)
            self._set_run_korean_font(run)

        # Save to bytes
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_from_draft_text(
        self,
        draft_text: str,
        case_title: str,
        citations: Optional[List[Dict[str, Any]]] = None,
        generated_at: Optional[datetime] = None
    ) -> bytes:
        """
        Generate DOCX from plain draft text (backward compatibility).

        Args:
            draft_text: Plain text draft content
            case_title: Case title for header
            citations: List of citation dictionaries
            generated_at: Generation timestamp

        Returns:
            bytes: DOCX file content
        """
        self.doc = Document()
        self._setup_page_format()
        self._setup_styles()

        # Title
        title = self.doc.add_heading("이혼 소송 준비서면 (초안)", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_korean_font(title)

        # Case info
        self.doc.add_paragraph()
        case_info = self.doc.add_paragraph()
        run = case_info.add_run(f"사건명: {case_title}")
        run.bold = True
        self._set_run_korean_font(run)

        timestamp = generated_at or datetime.now()
        p = self.doc.add_paragraph(f"생성일시: {timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
        self._set_paragraph_korean_font(p)

        # Draft content
        heading = self.doc.add_heading("본문", level=1)
        self._set_korean_font(heading)

        for paragraph_text in draft_text.split("\n\n"):
            if paragraph_text.strip():
                p = self.doc.add_paragraph(paragraph_text.strip())
                self._set_paragraph_korean_font(p)

        # Citations
        if citations:
            heading = self.doc.add_heading("인용 증거", level=1)
            self._set_korean_font(heading)

            for i, citation in enumerate(citations, 1):
                p = self.doc.add_paragraph()
                run = p.add_run(f"[증거 {i}] ")
                run.bold = True
                self._set_run_korean_font(run)

                run2 = p.add_run(f"(ID: {citation.get('evidence_id', 'N/A')})")
                self._set_run_korean_font(run2)

                labels = citation.get("labels", [])
                p2 = self.doc.add_paragraph(f"  - 분류: {', '.join(labels) if labels else 'N/A'}")
                self._set_paragraph_korean_font(p2)

                p3 = self.doc.add_paragraph(f"  - 내용: {citation.get('snippet', '')}")
                self._set_paragraph_korean_font(p3)

        # Disclaimer
        self.doc.add_paragraph()
        disclaimer = self.doc.add_paragraph(
            "※ 본 문서는 AI가 생성한 초안이며, "
            "변호사의 검토 및 수정이 필요합니다."
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in disclaimer.runs:
            run.italic = True
            self._set_run_korean_font(run)

        # Save to bytes
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _setup_page_format(self):
        """Configure A4 page size and Korean court margins."""
        section = self.doc.sections[0]
        section.page_width = Mm(self.PAPER_WIDTH_MM)
        section.page_height = Mm(self.PAPER_HEIGHT_MM)
        section.top_margin = Mm(self.MARGIN_TOP_MM)
        section.bottom_margin = Mm(self.MARGIN_BOTTOM_MM)
        section.left_margin = Mm(self.MARGIN_LEFT_MM)
        section.right_margin = Mm(self.MARGIN_RIGHT_MM)

    def _setup_styles(self):
        """Configure document styles with Korean fonts."""
        styles = self.doc.styles

        # Normal style
        normal = styles['Normal']
        normal.font.name = self.FONT_NAME_FALLBACK
        normal.font.size = Pt(self.FONT_SIZE_BODY)
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME_PRIMARY)

        # Configure paragraph format
        normal.paragraph_format.line_spacing = self.LINE_SPACING
        normal.paragraph_format.space_after = Pt(6)

    def _add_case_header(self, header: Dict[str, Any]):
        """Add case header section with court name, case number, parties."""
        if not header:
            return

        # Court name
        if header.get("court_name"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header["court_name"])
            run.bold = True
            run.font.size = Pt(self.FONT_SIZE_HEADING1)
            self._set_run_korean_font(run)

        # Case number
        if header.get("case_number"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"사건번호: {header['case_number']}")
            self._set_run_korean_font(run)

        # Parties
        parties = header.get("parties", {})
        if parties:
            self.doc.add_paragraph()
            if parties.get("plaintiff"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"원 고: {parties['plaintiff']}")
                self._set_run_korean_font(run)
            if parties.get("defendant"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"피 고: {parties['defendant']}")
                self._set_run_korean_font(run)

        self.doc.add_paragraph()

    def _add_document_title(self, document_type: str):
        """Add document type title (소장, 준비서면, etc.)."""
        title_text = self.DOCUMENT_TYPES.get(document_type, "준 비 서 면")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        self._set_run_korean_font(run)
        self.doc.add_paragraph()

    def _add_sections(self, sections: List[Dict[str, Any]]):
        """Add content sections (청구취지, 청구원인, etc.)."""
        sorted_sections = sorted(sections, key=lambda s: s.get("order", 0))

        for section in sorted_sections:
            # Section heading
            heading = self.doc.add_heading(section.get("title", ""), level=1)
            self._set_korean_font(heading)

            # Section content
            content = section.get("content", "")
            for paragraph_text in content.split("\n"):
                if paragraph_text.strip():
                    p = self.doc.add_paragraph(paragraph_text.strip())
                    self._set_paragraph_korean_font(p)

            self.doc.add_paragraph()

    def _add_citations(self, citations: List[Dict[str, Any]]):
        """Add evidence citations section."""
        if not citations:
            return

        heading = self.doc.add_heading("증거목록", level=1)
        self._set_korean_font(heading)

        for citation in citations:
            ref = citation.get("reference", "")
            desc = citation.get("description", "")
            p = self.doc.add_paragraph(f"{ref}: {desc}")
            self._set_paragraph_korean_font(p)

    def _add_signature(self, footer: Dict[str, Any]):
        """Add signature section with date and attorney name."""
        if not footer:
            return

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Date - right aligned
        if footer.get("date"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(footer["date"])
            self._set_run_korean_font(run)

        # Attorney name - right aligned
        if footer.get("attorney"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(footer["attorney"])
            run.bold = True
            self._set_run_korean_font(run)

            # Seal marker
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("(인)")
            run.font.size = Pt(10)
            self._set_run_korean_font(run)

    def _add_court_submission(self, court_name: str):
        """Add court submission line at the end."""
        if not court_name:
            return

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{court_name} 귀중")
        run.bold = True
        self._set_run_korean_font(run)

    def _set_korean_font(self, paragraph):
        """Set Korean font for a paragraph (used for headings)."""
        for run in paragraph.runs:
            self._set_run_korean_font(run)

    def _set_paragraph_korean_font(self, paragraph):
        """Set Korean font for all runs in a paragraph."""
        for run in paragraph.runs:
            self._set_run_korean_font(run)

    def _set_run_korean_font(self, run):
        """Set Korean font for a single run."""
        run.font.name = self.FONT_NAME_FALLBACK
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME_PRIMARY)


def generate_docx(
    content: Dict[str, Any],
    document_type: str = "brief"
) -> bytes:
    """
    Convenience function to generate DOCX from content.

    Args:
        content: Structured content dictionary
        document_type: Document type ("complaint", "motion", "brief", "response")

    Returns:
        bytes: DOCX file content
    """
    generator = DocxGenerator()
    return generator.generate_document(content, document_type)


def generate_docx_from_text(
    draft_text: str,
    case_title: str,
    citations: Optional[List[Dict[str, Any]]] = None,
    generated_at: Optional[datetime] = None
) -> bytes:
    """
    Convenience function to generate DOCX from plain text.

    Args:
        draft_text: Plain text draft content
        case_title: Case title
        citations: Optional list of citations
        generated_at: Optional generation timestamp

    Returns:
        bytes: DOCX file content
    """
    generator = DocxGenerator()
    return generator.generate_from_draft_text(
        draft_text, case_title, citations, generated_at
    )


def generate_docx_from_lines(
    lines: List[Dict[str, Any]],
    case_title: Optional[str] = None
) -> bytes:
    """
    Convenience function to generate DOCX from line-based template.

    Args:
        lines: List of line dictionaries with text and format info
        case_title: Optional case title

    Returns:
        bytes: DOCX file content
    """
    generator = DocxGenerator()
    return generator.generate_from_lines(lines, case_title)
