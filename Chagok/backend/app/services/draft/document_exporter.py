"""
Document Exporter - Generates DOCX and PDF files from draft content
Extracted from DraftService for better modularity (Issue #325)
"""

from typing import Tuple, List, Any
from io import BytesIO
from datetime import datetime
import os

from app.middleware import ValidationError

# Optional: python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentExporter:
    """
    Exports draft documents to DOCX and PDF formats.
    """

    def generate_docx(
        self,
        case: Any,
        draft_response: Any
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate DOCX file from draft response

        Args:
            case: Case object
            draft_response: Generated draft preview

        Returns:
            Tuple of (file_bytes, filename, content_type)
        """
        if not DOCX_AVAILABLE:
            raise ValidationError(
                "DOCX export is not available. "
                "Please install python-docx: pip install python-docx"
            )

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("이혼 소송 준비서면 (초안)", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case info
        doc.add_paragraph()
        case_info = doc.add_paragraph()
        case_info.add_run(f"사건명: {case.title}").bold = True
        doc.add_paragraph(f"생성일시: {draft_response.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")

        # Draft content
        doc.add_heading("본문", level=1)
        for paragraph_text in draft_response.draft_text.split("\n\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

        # Citations
        if draft_response.citations:
            doc.add_heading("인용 증거", level=1)
            for i, citation in enumerate(draft_response.citations, 1):
                p = doc.add_paragraph()
                p.add_run(f"[증거 {i}] ").bold = True
                p.add_run(f"(ID: {citation.evidence_id})")
                doc.add_paragraph(f"  - 분류: {', '.join(citation.labels) if citation.labels else 'N/A'}")
                doc.add_paragraph(f"  - 내용: {citation.snippet}")

        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "⚠️ 본 문서는 AI가 생성한 초안이며, "
            "변호사의 검토 및 수정이 필수입니다."
        ).italic = True

        # Save to BytesIO
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        # Generate filename
        safe_title = case.title.replace(" ", "_")[:30]
        filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return file_buffer, filename, content_type

    def generate_pdf(
        self,
        case: Any,
        draft_response: Any
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate PDF file from draft response with Korean font support

        Args:
            case: Case object
            draft_response: Generated draft preview

        Returns:
            Tuple of (file_bytes, filename, content_type)
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            )
            from reportlab.lib.units import inch, mm
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

            # Register Korean font
            korean_font_registered = self._register_korean_font(pdfmetrics, TTFont)
            font_name = "NotoSansKR" if korean_font_registered else "Helvetica"

            file_buffer = BytesIO()

            # A4 with proper margins for legal documents
            doc = SimpleDocTemplate(
                file_buffer,
                pagesize=A4,
                leftMargin=25 * mm,
                rightMargin=25 * mm,
                topMargin=25 * mm,
                bottomMargin=25 * mm
            )

            # Create custom styles with Korean font
            styles = getSampleStyleSheet()

            # Title style
            title_style = ParagraphStyle(
                'KoreanTitle',
                parent=styles['Title'],
                fontName=font_name,
                fontSize=18,
                alignment=TA_CENTER,
                spaceAfter=20
            )

            # Heading style
            heading_style = ParagraphStyle(
                'KoreanHeading',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=14,
                spaceBefore=15,
                spaceAfter=10
            )

            # Normal text style
            normal_style = ParagraphStyle(
                'KoreanNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                leading=16,
                alignment=TA_JUSTIFY,
                spaceAfter=8
            )

            # Citation style
            citation_style = ParagraphStyle(
                'Citation',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                leading=14,
                leftIndent=10,
                spaceAfter=6
            )

            # Disclaimer style
            disclaimer_style = ParagraphStyle(
                'Disclaimer',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                textColor=colors.grey,
                alignment=TA_CENTER,
                spaceBefore=20
            )

            story = []

            # === Document Header ===
            story.append(Paragraph("이혼 소송 준비서면", title_style))
            story.append(Paragraph("(초 안)", ParagraphStyle(
                'Subtitle',
                parent=normal_style,
                alignment=TA_CENTER,
                fontSize=12,
                spaceAfter=30
            )))

            # === Case Information Table ===
            case_data = [
                ["사 건 명", case.title],
                ["생성일시", draft_response.generated_at.strftime('%Y년 %m월 %d일 %H:%M')],
            ]
            case_table = Table(case_data, colWidths=[80, 350])
            case_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(case_table)
            story.append(Spacer(1, 0.4 * inch))

            # === Main Content ===
            story.append(Paragraph("본 문", heading_style))

            # Split content by paragraphs and add to story
            for paragraph_text in draft_response.draft_text.split("\n\n"):
                cleaned_text = paragraph_text.strip()
                if cleaned_text:
                    # Escape XML special characters
                    cleaned_text = (
                        cleaned_text
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(cleaned_text, normal_style))

            # === Citations Section ===
            if draft_response.citations:
                story.append(Spacer(1, 0.3 * inch))
                story.append(Paragraph("인용 증거", heading_style))

                for i, citation in enumerate(draft_response.citations, 1):
                    # Citation header
                    labels_str = ", ".join(citation.labels) if citation.labels else "N/A"
                    citation_header = f"<b>[증거 {i}]</b> (ID: {citation.evidence_id})"
                    story.append(Paragraph(citation_header, citation_style))

                    # Citation details
                    story.append(Paragraph(f"분류: {labels_str}", citation_style))

                    # Citation snippet (escape special characters)
                    snippet = (
                        citation.snippet
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(f"내용: {snippet}", citation_style))
                    story.append(Spacer(1, 0.1 * inch))

            # === Disclaimer ===
            story.append(Spacer(1, 0.5 * inch))
            story.append(Paragraph(
                "⚠ 본 문서는 AI가 생성한 초안이며, 변호사의 검토 및 수정이 필수입니다.",
                disclaimer_style
            ))

            # Build PDF
            doc.build(story)
            file_buffer.seek(0)

            safe_title = case.title.replace(" ", "_")[:30]
            filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.pdf"
            content_type = "application/pdf"

            return file_buffer, filename, content_type

        except ImportError:
            raise ValidationError(
                "PDF export is not available. "
                "Please install reportlab: pip install reportlab. "
                "Alternatively, use DOCX format."
            )

    def _register_korean_font(self, pdfmetrics, TTFont) -> bool:
        """
        Register Korean font for PDF generation

        Args:
            pdfmetrics: reportlab pdfmetrics module
            TTFont: reportlab TTFont class

        Returns:
            bool: True if Korean font was registered, False otherwise
        """
        # Font search paths
        font_candidates = [
            # Bundled font (if exists in project)
            os.path.join(os.path.dirname(__file__), "..", "..", "fonts", "NotoSansKR-Regular.ttf"),
            # macOS system fonts
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
            "/Library/Fonts/NotoSansKR-Regular.ttf",
            # Linux system fonts
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            # Windows system fonts
            "C:/Windows/Fonts/malgun.ttf",
            "C:/Windows/Fonts/NotoSansKR-Regular.ttf",
        ]

        for font_path in font_candidates:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("NotoSansKR", font_path))
                    return True
                except Exception:
                    continue

        # No Korean font found - will use default font
        return False

    def generate_docx_from_lines(
        self,
        case: Any,
        lines: List[dict]
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate DOCX directly from line-based JSON

        Args:
            case: Case object
            lines: List of line dicts with format info

        Returns:
            Tuple of (file_buffer, filename, content_type)
        """
        if not DOCX_AVAILABLE:
            raise ValidationError("DOCX export is not available. Please install python-docx.")

        doc = Document()

        for line in lines:
            text = line.get("text", "")
            fmt = line.get("format", {})

            # Skip empty lines (just add spacing)
            if not text:
                doc.add_paragraph("")
                continue

            para = doc.add_paragraph()

            # Apply alignment
            align = fmt.get("align", "left")
            if align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add text with formatting
            run = para.add_run(text)

            if fmt.get("bold", False):
                run.bold = True

            font_size = fmt.get("font_size", 12)
            run.font.size = Pt(font_size)

            # Handle indent as paragraph format
            indent = fmt.get("indent", 0)
            if indent > 0:
                para.paragraph_format.left_indent = Pt(indent * 6)  # 6pt per char approx

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"draft_{case.title[:20]}_{timestamp}.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return buffer, filename, content_type

    def generate_pdf_from_lines(
        self,
        case: Any,
        lines: List[dict]
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate PDF directly from line-based JSON

        Args:
            case: Case object
            lines: List of line dicts with format info

        Returns:
            Tuple of (file_buffer, filename, content_type)
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ValidationError("PDF export is not available. Please install reportlab.")

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)

        # Register Korean font
        self._register_korean_font(pdfmetrics, TTFont)

        styles = getSampleStyleSheet()

        # Create custom styles
        style_left = ParagraphStyle('Left', parent=styles['Normal'], alignment=TA_LEFT)
        style_center = ParagraphStyle('Center', parent=styles['Normal'], alignment=TA_CENTER)
        style_right = ParagraphStyle('Right', parent=styles['Normal'], alignment=TA_RIGHT)

        story = []

        for line in lines:
            text = line.get("text", "")
            fmt = line.get("format", {})

            if not text:
                story.append(Spacer(1, 12))
                continue

            # Choose style based on alignment
            align = fmt.get("align", "left")
            if align == "center":
                style = style_center
            elif align == "right":
                style = style_right
            else:
                style = style_left

            # Apply bold if needed
            if fmt.get("bold", False):
                text = f"<b>{text}</b>"

            # Apply indent
            indent = fmt.get("indent", 0)
            if indent > 0:
                text = "&nbsp;" * indent + text

            para = Paragraph(text, style)
            story.append(para)

            # Add spacing
            spacing = fmt.get("spacing_after", 0)
            if spacing > 0:
                story.append(Spacer(1, spacing * 12))

        doc.build(story)
        buffer.seek(0)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"draft_{case.title[:20]}_{timestamp}.pdf"
        content_type = "application/pdf"

        return buffer, filename, content_type
