"""
Draft Exporter Module
DOCX/PDF 내보내기 기능
"""

from typing import Tuple
from io import BytesIO

from app.db.schemas import DraftPreviewResponse
from app.middleware import ValidationError

# Optional: python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Pt, Inches  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_docx(
    case,
    draft_response: DraftPreviewResponse
) -> Tuple[BytesIO, str, str]:
    """
    Generate DOCX file from draft response

    Args:
        case: Case object
        draft_response: Generated draft preview

    Returns:
        Tuple of (file_bytes, filename, content_type)
    """
    if not DOCX_AVAILABLE:
        raise ValidationError(
            "DOCX export is not available. "
            "Please install python-docx: pip install python-docx"
        )

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading("이혼 소송 준비서면 (초안)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case info
    doc.add_paragraph()
    case_info = doc.add_paragraph()
    case_info.add_run(f"사건명: {case.title}").bold = True
    doc.add_paragraph(f"생성일시: {draft_response.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")

    # Draft content
    doc.add_heading("본문", level=1)
    for paragraph_text in draft_response.draft_text.split("\n\n"):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text.strip())

    # Citations
    if draft_response.citations:
        doc.add_heading("인용 증거", level=1)
        for i, citation in enumerate(draft_response.citations, 1):
            p = doc.add_paragraph()
            p.add_run(f"[증거 {i}] ").bold = True
            p.add_run(f"(ID: {citation.evidence_id})")
            doc.add_paragraph(f"  - 분류: {', '.join(citation.labels) if citation.labels else 'N/A'}")
            doc.add_paragraph(f"  - 내용: {citation.snippet}")

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run(
        "⚠️ 본 문서는 AI가 생성한 초안이며, "
        "변호사의 검토 및 수정이 필수입니다."
    ).italic = True

    # Save to BytesIO
    file_buffer = BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)

    # Generate filename
    safe_title = case.title.replace(" ", "_")[:30]
    filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return file_buffer, filename, content_type


def generate_pdf(
    case,
    draft_response: DraftPreviewResponse
) -> Tuple[BytesIO, str, str]:
    """
    Generate PDF file from draft response

    Note: For simplicity, we generate a basic text-based PDF.
    For production, consider using reportlab or weasyprint.

    Args:
        case: Case object
        draft_response: Generated draft preview

    Returns:
        Tuple of (file_bytes, filename, content_type)
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase.ttfonts import TTFont  # noqa: F401

        file_buffer = BytesIO()
        doc = SimpleDocTemplate(file_buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph("이혼 소송 준비서면 (초안)", styles['Title']))
        story.append(Spacer(1, 0.5 * inch))

        # Case info
        story.append(Paragraph(f"<b>사건명:</b> {case.title}", styles['Normal']))
        story.append(Paragraph(
            f"<b>생성일시:</b> {draft_response.generated_at.strftime('%Y-%m-%d %H:%M:%S')}",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.3 * inch))

        # Draft content
        story.append(Paragraph("<b>본문</b>", styles['Heading2']))
        for paragraph_text in draft_response.draft_text.split("\n\n"):
            if paragraph_text.strip():
                story.append(Paragraph(paragraph_text.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1 * inch))

        # Build PDF
        doc.build(story)
        file_buffer.seek(0)

        safe_title = case.title.replace(" ", "_")[:30]
        filename = f"draft_{safe_title}_{draft_response.generated_at.strftime('%Y%m%d')}.pdf"
        content_type = "application/pdf"

        return file_buffer, filename, content_type

    except ImportError:
        raise ValidationError(
            "PDF export is not available. "
            "Please install reportlab: pip install reportlab. "
            "Alternatively, use DOCX format."
        )
